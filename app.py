import streamlit as st
import pandas as pd
import openai
import random
from docx import Document
import zipfile
import io

# OpenAI API key setup (use secrets or environment variable for security)
openai.api_key = st.secrets["openai"]["api_key"]

def generate_question(case_vignette, question_type):
    prompt = f"""Given the following pediatric case vignette:
{case_vignette}

Generate a USMLE/NBME style multiple choice question for: {question_type}.
The question must be formatted in multiple choice style with exactly four answer options labeled A, B, C, and D.
Clearly indicate the correct answer.
Also, provide a detailed explanation for why that answer is correct.
Format your response using the following sections and delimiters exactly:

---BEGIN QUESTION---
Question:
<Your multiple choice question including answer options>
---END QUESTION---

---BEGIN ANSWER KEY---
Answer Key:
<The correct answer letter (e.g., A, B, C, or D)>
---END ANSWER KEY---

---BEGIN EXPLANATION---
Explanation:
<The detailed explanation for the correct answer>
---END EXPLANATION---
"""
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are an expert in pediatric medical education."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=500,
        temperature=0.7
    )
    return response['choices'][0]['message']['content'].strip()


def create_case_vignette(row):
    # Customize how the case vignette is built using the diagnoses columns.
    vignette = (
        f"A pediatric patient presents with {row['Top Diagnosis']}. "
        f"Differential diagnoses include {row['Diagnosis 1']}, {row['Diagnosis 2']}, and {row['Diagnosis 3']}."
    )
    return vignette

st.title("Pediatric Shelf Examination Generator")

uploaded_file = st.file_uploader("Upload CSV file", type=["csv"])
if uploaded_file:
    df = pd.read_csv(uploaded_file)
    st.write("Data Preview:", df.head())

    # Group by the email column
    grouped = df.groupby("Please enter email.")
    
    # Create a zip buffer to store all the zip files for download if desired
    zip_all_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_all_buffer, "w", zipfile.ZIP_DEFLATED) as main_zip:
        # Process each group
        for email, group in grouped:
            doc = Document()
            doc.add_heading(f"Pediatric Shelf Examination for {email}", 0)
            
            # Track case number within the group for clarity.
            case_number = 1
            for idx, row in group.iterrows():
                vignette = create_case_vignette(row)
                # Randomly choose anchoring question type (or set logic)
                question_type = random.choice(["next best step in management", "most likely diagnosis"])
                
                # Generate two questions per entry
                question1 = generate_question(vignette, question_type)
                question2 = generate_question(vignette, question_type)
                
                # Add a new section for each case
                doc.add_heading(f"Case {case_number}", level=1)
                doc.add_paragraph("Case Vignette:")
                doc.add_paragraph(vignette)
                doc.add_paragraph(f"Question 1 ({question_type.title()}):")
                doc.add_paragraph(question1)
                doc.add_paragraph(f"Question 2 ({question_type.title()}):")
                doc.add_paragraph(question2)
                case_number += 1

            # Save the Word document into a bytes buffer
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            # Write the document into the main zip file with the email as the filename
            main_zip.writestr(f"{email}.docx", doc_buffer.getvalue())
    
    zip_all_buffer.seek(0)
    st.download_button(
        label="Download All Exams as ZIP",
        data=zip_all_buffer,
        file_name="exams.zip",
        mime="application/zip"
    )
